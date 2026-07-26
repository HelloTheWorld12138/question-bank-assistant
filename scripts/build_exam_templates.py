"""Build the bundled Pandoc reference DOCX files.

Design baseline: compact_reference_guide with a named ``physics_exam_a4``
override for Chinese printed examination papers.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT / "templates"
INK = RGBColor(0x17, 0x21, 0x2B)
MUTED = RGBColor(0x5F, 0x6F, 0x82)
# The delivery target is Chinese Windows 10/11. SimSun and SimHei are stable
# Word/WPS choices there; OfficeCLI supplies a readable fallback in previews.
BODY_CJK_FONT = "SimSun"
HEADING_CJK_FONT = "SimHei"
BODY_FONT = "Times New Roman"
HEADING_FONT = "Arial"


TEMPLATE_SPECS = {
    "a4_single.docx": {
        "label": "A4 单栏练习",
        "columns": 1,
        "margins": (2.0, 2.0, 2.0, 2.0),
    },
    "a4_double.docx": {
        "label": "A4 双栏练习",
        "columns": 2,
        "margins": (1.6, 1.6, 1.6, 1.6),
    },
    "formal_exam.docx": {
        "label": "正式考试卷",
        "columns": 1,
        "margins": (2.0, 2.2, 2.0, 2.2),
    },
}


def set_run_fonts(run, latin: str, east_asia: str) -> None:
    run.font.name = latin
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:hint"), "eastAsia")
    language = properties.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        properties.append(language)
    language.set(qn("w:val"), "zh-CN")
    language.set(qn("w:eastAsia"), "zh-CN")


def configure_style(
    style,
    *,
    latin: str,
    east_asia: str,
    size: float,
    bold: bool = False,
    color: RGBColor = INK,
    align=None,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.25,
    keep_next: bool = False,
) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:hint"), "eastAsia")
    properties = style.element.get_or_add_rPr()
    language = properties.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        properties.append(language)
    language.set(qn("w:val"), "zh-CN")
    language.set(qn("w:eastAsia"), "zh-CN")
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing
    paragraph.keep_with_next = keep_next
    if align is not None:
        paragraph.alignment = align
    # LibreOffice's built-in Title style can carry a colored bottom border.
    # Reference documents should remain neutral, so strip any inherited border.
    paragraph_properties = style.element.get_or_add_pPr()
    border = paragraph_properties.find(qn("w:pBdr"))
    if border is not None:
        paragraph_properties.remove(border)


def set_columns(section, count: int) -> None:
    section_properties = section._sectPr
    columns = section_properties.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        section_properties.append(columns)
    columns.set(qn("w:num"), str(count))
    columns.set(qn("w:space"), "425" if count == 2 else "720")
    columns.set(qn("w:equalWidth"), "1")


def append_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    run = paragraph.add_run()
    set_run_fonts(run, BODY_FONT, BODY_CJK_FONT)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_text, separate, value, end])


def configure_document_styles(document: Document) -> None:
    styles = document.styles
    configure_style(
        styles["Normal"],
        latin=BODY_FONT,
        east_asia=BODY_CJK_FONT,
        size=10.5,
        after=0,
        line_spacing=1.25,
    )
    configure_style(
        styles["Title"],
        latin=HEADING_FONT,
        east_asia=HEADING_CJK_FONT,
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=6,
        line_spacing=1.0,
        keep_next=True,
    )
    configure_style(
        styles["Subtitle"],
        latin=BODY_FONT,
        east_asia=BODY_CJK_FONT,
        size=10.5,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line_spacing=1.0,
        keep_next=True,
    )
    for name, size, before, after in (
        ("Heading 1", 12, 10, 4),
        ("Heading 2", 11, 8, 3),
        ("Heading 3", 10.5, 6, 2),
    ):
        configure_style(
            styles[name],
            latin=HEADING_FONT,
            east_asia=HEADING_CJK_FONT,
            size=size,
            bold=True,
            before=before,
            after=after,
            line_spacing=1.0,
            keep_next=True,
        )

    if "Question Number" not in styles:
        question_number = styles.add_style("Question Number", WD_STYLE_TYPE.PARAGRAPH)
    else:
        question_number = styles["Question Number"]
    configure_style(
        question_number,
        latin=BODY_FONT,
        east_asia=BODY_CJK_FONT,
        size=10.5,
        bold=True,
        after=2,
        line_spacing=1.25,
        keep_next=True,
    )

    # Pandoc creates syntax-token character styles based on VerbatimChar.
    # Providing the base style prevents dangling style references in the
    # resulting DOCX even when the paper contains no source-code block.
    if "Verbatim Char" not in styles:
        verbatim = styles.add_style("Verbatim Char", WD_STYLE_TYPE.CHARACTER)
    else:
        verbatim = styles["Verbatim Char"]
    verbatim.font.name = "Consolas"
    verbatim.font.size = Pt(9.5)
    verbatim_properties = verbatim.element.get_or_add_rPr()
    verbatim_fonts = verbatim_properties.get_or_add_rFonts()
    verbatim_fonts.set(qn("w:ascii"), "Consolas")
    verbatim_fonts.set(qn("w:hAnsi"), "Consolas")
    verbatim_fonts.set(qn("w:eastAsia"), BODY_CJK_FONT)


def configure_section(document: Document, spec: dict) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    top, right, bottom, left = spec["margins"]
    section.top_margin = Cm(top)
    section.right_margin = Cm(right)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)
    set_columns(section, spec["columns"])

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run(f"高中物理题库助手 · {spec['label']}")
    set_run_fonts(header_run, BODY_FONT, BODY_CJK_FONT)
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = MUTED

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    prefix = footer_paragraph.add_run("第 ")
    set_run_fonts(prefix, BODY_FONT, BODY_CJK_FONT)
    prefix.font.size = Pt(8.5)
    append_field(footer_paragraph, "PAGE")
    suffix = footer_paragraph.add_run(" 页")
    set_run_fonts(suffix, BODY_FONT, BODY_CJK_FONT)
    suffix.font.size = Pt(8.5)


def add_preview_content(document: Document, spec: dict) -> None:
    document.add_paragraph("高中物理试卷（样式模板）", style="Title")
    document.add_paragraph(
        "班级：________　姓名：________　考试时间：45 分钟　满分：100 分",
        style="Subtitle",
    )
    document.add_paragraph("一、选择题", style="Heading 1")
    document.add_paragraph("1.【LXJD0001】", style="Question Number")
    document.add_paragraph(
        "如图所示，质量为 m 的物体在恒力 F 作用下做匀加速直线运动。"
        "若物体的加速度为 a，请写出三者之间的关系式。"
    )
    equation = document.add_paragraph()
    equation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_run = equation.add_run("F = ma")
    set_run_fonts(equation_run, "Cambria Math", BODY_CJK_FONT)
    equation_run.font.size = Pt(12)
    document.add_paragraph("二、计算题", style="Heading 1")
    document.add_paragraph("2.【LXCX0001】", style="Question Number")
    document.add_paragraph(
        "请写出必要的文字说明、方程式和演算步骤。答案中的物理量应标明单位。"
    )
    if spec["columns"] == 2:
        for index in range(3, 9):
            document.add_paragraph(f"{index}. 双栏版式示例题目，用于检查栏宽、换栏和行距。")


def build_template(filename: str, spec: dict) -> Path:
    document = Document()
    configure_document_styles(document)
    configure_section(document, spec)
    add_preview_content(document, spec)
    output = OUTPUT_DIR / filename
    document.save(output)
    return output


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for filename, spec in TEMPLATE_SPECS.items():
        print(build_template(filename, spec))


if __name__ == "__main__":
    main()
